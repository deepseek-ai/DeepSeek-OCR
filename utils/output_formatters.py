"""
Output Format Converters
Support for JSON, HTML, DOCX, and CSV output formats
"""

import json
import re
from typing import Dict, List, Any, Optional
from pathlib import Path
import io


class OutputFormatter:
    """Base class for output formatters"""

    @staticmethod
    def extract_elements(text: str, matches: List[tuple]) -> List[Dict[str, Any]]:
        """Extract structured elements from OCR output"""
        elements = []

        for match in matches:
            try:
                full_match, label_type, coords = match
                coords_list = eval(coords)

                elements.append({
                    "type": label_type,
                    "coordinates": coords_list,
                    "raw": full_match
                })
            except:
                continue

        return elements

    @staticmethod
    def extract_tables(text: str) -> List[Dict[str, Any]]:
        """Extract tables from markdown text"""
        tables = []
        table_pattern = r'\|(.+)\|[\r\n]+\|[\s\-:]+\|(.*?)(?=\n\n|\n\|[^\-]|\Z)'

        for match in re.finditer(table_pattern, text, re.DOTALL):
            table_text = match.group(0)
            lines = [line.strip() for line in table_text.split('\n') if line.strip()]

            if len(lines) >= 2:
                headers = [cell.strip() for cell in lines[0].split('|') if cell.strip()]
                rows = []

                for line in lines[2:]:  # Skip header and separator
                    cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                    if cells:
                        rows.append(cells)

                tables.append({
                    "headers": headers,
                    "rows": rows,
                    "raw": table_text
                })

        return tables


class JSONFormatter(OutputFormatter):
    """Format OCR output as JSON"""

    @staticmethod
    def format(text: str, matches: List[tuple], image_width: int = 0, image_height: int = 0,
               metadata: Optional[Dict[str, Any]] = None) -> str:
        """Convert OCR output to JSON format"""

        # Extract structured data
        elements = OutputFormatter.extract_elements(text, matches)
        tables = OutputFormatter.extract_tables(text)

        # Clean text (remove grounding tags)
        clean_text = text
        for match in matches:
            clean_text = clean_text.replace(match[0], '')

        clean_text = clean_text.replace('\\coloneqq', ':=').replace('\\eqqcolon', '=:')
        clean_text = re.sub(r'\n{3,}', '\n\n', clean_text)

        output = {
            "text": clean_text.strip(),
            "elements": elements,
            "tables": tables,
            "metadata": metadata or {},
            "dimensions": {
                "width": image_width,
                "height": image_height
            }
        }

        return json.dumps(output, indent=2, ensure_ascii=False)


class HTMLFormatter(OutputFormatter):
    """Format OCR output as HTML"""

    @staticmethod
    def format(text: str, matches: List[tuple], metadata: Optional[Dict[str, Any]] = None) -> str:
        """Convert OCR output to styled HTML"""

        # Clean text
        clean_text = text
        for match in matches:
            clean_text = clean_text.replace(match[0], '')

        clean_text = clean_text.replace('\\coloneqq', ':=').replace('\\eqqcolon', '=:')

        # Convert markdown to HTML-like structure
        html_text = clean_text

        # Headers
        html_text = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html_text, flags=re.MULTILINE)
        html_text = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html_text, flags=re.MULTILINE)
        html_text = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html_text, flags=re.MULTILINE)

        # Bold and italic
        html_text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html_text)
        html_text = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html_text)

        # Links
        html_text = re.sub(r'\[(.+?)\]\((.+?)\)', r'<a href="\2">\1</a>', html_text)

        # Code blocks
        html_text = re.sub(r'```(.+?)```', r'<pre><code>\1</code></pre>', html_text, flags=re.DOTALL)
        html_text = re.sub(r'`(.+?)`', r'<code>\1</code>', html_text)

        # Lists
        html_text = re.sub(r'^\- (.+)$', r'<li>\1</li>', html_text, flags=re.MULTILINE)
        html_text = re.sub(r'((?:<li>.+</li>\n)+)', r'<ul>\n\1</ul>\n', html_text)

        # Paragraphs
        paragraphs = html_text.split('\n\n')
        html_paragraphs = []
        for p in paragraphs:
            p = p.strip()
            if p and not p.startswith('<'):
                html_paragraphs.append(f'<p>{p}</p>')
            elif p:
                html_paragraphs.append(p)

        html_body = '\n'.join(html_paragraphs)

        # Build complete HTML
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>OCR Result</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f5f5f5;
        }}
        .container {{
            background-color: white;
            padding: 30px;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        h1 {{ color: #2c3e50; border-bottom: 3px solid #3498db; padding-bottom: 10px; }}
        h2 {{ color: #34495e; border-bottom: 2px solid #95a5a6; padding-bottom: 8px; }}
        h3 {{ color: #555; }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }}
        th {{
            background-color: #3498db;
            color: white;
            font-weight: bold;
        }}
        tr:nth-child(even) {{
            background-color: #f2f2f2;
        }}
        code {{
            background-color: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
        }}
        pre {{
            background-color: #2d2d2d;
            color: #f8f8f2;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
        }}
        pre code {{
            background-color: transparent;
            color: inherit;
        }}
        .metadata {{
            background-color: #ecf0f1;
            padding: 15px;
            border-radius: 5px;
            margin-bottom: 20px;
            font-size: 0.9em;
        }}
        ul {{ padding-left: 30px; }}
        li {{ margin: 8px 0; }}
    </style>
</head>
<body>
    <div class="container">
        {f'<div class="metadata"><strong>Metadata:</strong> {metadata}</div>' if metadata else ''}
        {html_body}
    </div>
</body>
</html>"""

        return html


class DOCXFormatter(OutputFormatter):
    """Format OCR output as DOCX"""

    @staticmethod
    def format(text: str, matches: List[tuple], output_path: str,
               metadata: Optional[Dict[str, Any]] = None) -> str:
        """Convert OCR output to DOCX format"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            # Clean text
            clean_text = text
            for match in matches:
                clean_text = clean_text.replace(match[0], '')

            clean_text = clean_text.replace('\\coloneqq', ':=').replace('\\eqqcolon', '=:')

            doc = Document()

            # Add metadata if provided
            if metadata:
                doc.core_properties.title = metadata.get('title', 'OCR Result')
                doc.core_properties.author = metadata.get('author', 'DeepSeek-OCR')

            # Process text line by line
            for line in clean_text.split('\n'):
                line = line.strip()
                if not line:
                    continue

                # Headers
                if line.startswith('### '):
                    p = doc.add_heading(line[4:], level=3)
                elif line.startswith('## '):
                    p = doc.add_heading(line[3:], level=2)
                elif line.startswith('# '):
                    p = doc.add_heading(line[2:], level=1)
                # Lists
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif re.match(r'^\d+\.\s', line):
                    doc.add_paragraph(re.sub(r'^\d+\.\s', '', line), style='List Number')
                # Tables (basic support)
                elif line.startswith('|') and '|' in line[1:]:
                    # Skip for now - tables need special handling
                    continue
                # Regular paragraph
                else:
                    doc.add_paragraph(line)

            doc.save(output_path)
            return output_path

        except ImportError:
            raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")


class CSVFormatter(OutputFormatter):
    """Format OCR table output as CSV"""

    @staticmethod
    def format(text: str, output_path: str) -> str:
        """Extract tables and convert to CSV format"""
        import csv

        tables = OutputFormatter.extract_tables(text)

        if not tables:
            # No tables found, create simple CSV from text
            with open(output_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerow(['Text'])
                for line in text.split('\n'):
                    if line.strip():
                        writer.writerow([line.strip()])
            return output_path

        # Export first table (or combine multiple)
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)

            for i, table in enumerate(tables):
                if i > 0:
                    writer.writerow([])  # Empty row between tables
                    writer.writerow([f'--- Table {i + 1} ---'])

                writer.writerow(table['headers'])
                writer.writerows(table['rows'])

        return output_path


class ExcelFormatter(OutputFormatter):
    """Format OCR table output as Excel"""

    @staticmethod
    def format(text: str, output_path: str) -> str:
        """Extract tables and convert to Excel format"""
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment

            tables = OutputFormatter.extract_tables(text)

            wb = openpyxl.Workbook()
            wb.remove(wb.active)  # Remove default sheet

            if not tables:
                # No tables found, create simple sheet from text
                ws = wb.create_sheet("Text")
                ws.append(['Text'])
                for line in text.split('\n'):
                    if line.strip():
                        ws.append([line.strip()])
            else:
                # Create sheet for each table
                for i, table in enumerate(tables):
                    ws = wb.create_sheet(f"Table_{i + 1}")

                    # Headers
                    ws.append(table['headers'])

                    # Style headers
                    for cell in ws[1]:
                        cell.font = Font(bold=True, color="FFFFFF")
                        cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
                        cell.alignment = Alignment(horizontal='center')

                    # Data rows
                    for row in table['rows']:
                        ws.append(row)

                    # Auto-adjust column widths
                    for column in ws.columns:
                        max_length = 0
                        column_letter = column[0].column_letter
                        for cell in column:
                            try:
                                if len(str(cell.value)) > max_length:
                                    max_length = len(cell.value)
                            except:
                                pass
                        adjusted_width = min(max_length + 2, 50)
                        ws.column_dimensions[column_letter].width = adjusted_width

            wb.save(output_path)
            return output_path

        except ImportError:
            raise ImportError("openpyxl is required for Excel export. Install with: pip install openpyxl")
